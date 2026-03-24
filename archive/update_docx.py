"""Update API-CONTRACT.docx with current codebase state."""
from docx import Document
from copy import deepcopy

doc = Document("C:/claude_workspaces/hatchloom/hatchloom-delta/API-CONTRACT.docx")

# ============================================================
# 1. Auth token table (Table 0): Add parent token row
# ============================================================
auth_table = doc.tables[0]
new_row = auth_table.add_row()
new_row.cells[0].text = "`test-parent-token`"
new_row.cells[1].text = "14"
new_row.cells[2].text = "`parent`"
new_row.cells[3].text = "1"

# Copy formatting from existing data rows
for i, cell in enumerate(new_row.cells):
    src_cell = auth_table.rows[3].cells[i]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            src_run = src_cell.paragraphs[0].runs[0] if src_cell.paragraphs[0].runs else None
            if src_run:
                run.font.size = src_run.font.size
                run.font.name = src_run.font.name

# ============================================================
# 2. "Three hardcoded tokens" -> "Four hardcoded tokens"
# ============================================================
for p in doc.paragraphs:
    if "Three hardcoded tokens" in p.text:
        for run in p.runs:
            if "Three" in run.text:
                run.text = run.text.replace("Three", "Four")

# ============================================================
# 3. Auth behavior bullet: update role access description
# ============================================================
for p in doc.paragraphs:
    if "Valid token mapping to a user whose role is not" in p.text:
        for run in p.runs:
            run.text = ""
        p.runs[0].text = (
            "Valid token mapping to a user with role school_admin or school_teacher "
            "grants full read/write access. Students (role=student) have read-only "
            "access scoped to their own data. Parents (role=parent) have read-only "
            "access scoped to their linked child's data via the parent_of column. "
            "Students and parents receive 403 Forbidden on write attempts (POST, PUT, PATCH, DELETE)."
        )
        break

# ============================================================
# 4. Grade filter in enrolments query params (Table 5, row 3)
# ============================================================
grade_table = doc.tables[5]
grade_cell = grade_table.rows[3].cells[3]
for paragraph in grade_cell.paragraphs:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = "Filter students by grade level (8-12). Matches the grade column on the users table."

# ============================================================
# 5. Grade note paragraph: update
# ============================================================
for p in doc.paragraphs:
    if "The grade filter is accepted but has no effect because the users table does not yet have a grade column" in p.text:
        for run in p.runs:
            run.text = ""
        p.runs[0].text = (
            "Note: The student_id filter is used internally by the Experience Service "
            "to look up individual students within an experience context."
        )
        break

# ============================================================
# 6. Credential dependency note: remove "grade is always null"
# ============================================================
for p in doc.paragraphs:
    if "The grade field is always null until the users table is extended with a grade column" in p.text:
        new_text = p.text.replace(
            " The grade field is always null until the users table is extended with a grade column.",
            " Student grades (8-12) are now seeded in the users table.",
        )
        for run in p.runs:
            run.text = ""
        p.runs[0].text = new_text
        break

# ============================================================
# 7. Quebec roles: add "parent" to the role list
# ============================================================
for p in doc.paragraphs:
    if "user_id, role (one of school_admin, school_teacher, student), and school_id" in p.text:
        for run in p.runs:
            if "school_teacher, student)" in run.text:
                run.text = run.text.replace(
                    "school_teacher, student)", "school_teacher, student, parent)"
                )
        break

# ============================================================
# 8. Add "Getting Started" + "Seeded Test Data" before Auth
# ============================================================
auth_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Authentication" and p.style.name == "Heading 2":
        auth_idx = i
        break

if auth_idx:
    items = [
        ("Heading 2", "Getting Started"),
        ("Normal", "Clone the repo and start all services with Docker:"),
        ("Normal", "git clone <repo-url> hatchloom-delta && cd hatchloom-delta\ndocker compose up --build -d"),
        ("Normal", "Wait ~15 seconds, then verify all services are healthy:"),
        ("Normal", "curl http://localhost:8001/api/school/dashboard/health\ncurl http://localhost:8002/api/school/experiences/health\ncurl http://localhost:8003/api/school/enrolments/health"),
        ("Normal", "Migrations and seeding happen automatically -- no extra commands needed."),
        ("Heading 2", "Seeded Test Data"),
        ("Normal", "All services are pre-seeded with:"),
        ("List Bullet", "1 school: Ridgewood Academy (id=1)"),
        ("List Bullet", "14 users: 1 admin (id=1), 2 teachers (id=2,3), 10 students (id=4-13, grades 8-12), 1 parent (id=14, linked to student id=4)"),
        ("List Bullet", "2 experiences: Business Foundations, Tech Explorers (with 5 courses total)"),
        ("List Bullet", "3 cohorts: Cohort A (active, 6 students), Cohort B (not_started), Cohort C (active, 2 students)"),
        ("List Bullet", "5 mock courses: IDs 1-5 (hardcoded via MockCourseDataProvider)"),
        ("List Bullet", "Student grades: 8-12 distributed across students (Student 1 / id=4 is Grade 10)"),
        ("List Bullet", "Sample credential data: 2 earned + 1 in-progress per student (mock)"),
        ("Normal", "Tokens are forwarded: When Dashboard Service calls Experience/Enrolment, it forwards your bearer token. One token works across the whole chain."),
    ]
    for style_name, text in reversed(items):
        doc.paragraphs[auth_idx].insert_paragraph_before(text, style_name)

# ============================================================
# 9. Add Data Ownership + Response Format before Integration Requests
# ============================================================
integ_idx = None
for i, p in enumerate(doc.paragraphs):
    if "Integration Requests" in p.text and p.style.name == "Heading 2":
        integ_idx = i
        break

if integ_idx:
    sections = [
        ("Heading 2", "Response Format"),
        ("Normal", "All endpoints use a consistent response envelope:"),
        ("Normal", 'Success (list): { "data": [...], "meta": { "current_page": 1, "last_page": 1, "per_page": 15, "total": 3 } }'),
        ("Normal", 'Success (single): { "id": 1, "name": "Cohort A", "status": "active", ... }'),
        ("Normal", 'Error: { "error": true, "message": "Human-readable message", "code": "MACHINE_READABLE_CODE" }'),
        ("Heading 2", "Data Ownership -- What Comes From Where"),
        ("Normal", "Delta owns school administration data: experiences, cohorts, enrolments, and the admin dashboard. Delta does NOT own student activity, profile, gamification, messaging, or venture data."),
        ("Heading 3", "Data Delta provides"),
        ("List Bullet", "Student enrolment status -- GET /api/school/enrolments"),
        ("List Bullet", "Student detail (enrolment + dashboard context) -- GET /enrolments/students/{id}, GET /dashboard/students/{id}"),
        ("List Bullet", "Experience list, detail, course contents, statistics -- GET /api/school/experiences/*"),
        ("List Bullet", "Cohort list and detail -- GET /api/school/cohorts, /cohorts/{id}"),
        ("List Bullet", "Enrolment statistics with warnings -- GET /api/school/enrolments/statistics"),
        ("List Bullet", "CSV exports -- GET /enrolments/export, /experiences/{id}/students/export"),
        ("List Bullet", "Credential data (mock) and curriculum mapping (mock) -- embedded in student detail endpoints"),
        ("Heading 3", "Data Delta does NOT provide"),
        ("List Bullet", "User accounts, login, JWT auth -- Team Quebec (Auth). Delta uses mock bearer tokens."),
        ("List Bullet", "Parent-child relationships (multi-child) -- Team Quebec. Delta has single parent_of column. Multi-child requires a pivot table."),
        ("List Bullet", "Student profile (handle, avatar, bio, phone) -- Team Quebec. Delta only has name/email/role/school_id/grade."),
        ("List Bullet", "Course catalogue and course content -- Team Papa (Explore). Delta has MockCourseDataProvider with 5 stub courses."),
        ("List Bullet", "Per-student course progress -- Team Papa. Delta has no progress tracking."),
        ("List Bullet", "Credential engine (real credentials) -- Karl (Role B). Delta uses MockCredentialDataProvider."),
        ("List Bullet", "Gamification (streaks, XP, badges) -- TBD (not assigned to any team)."),
        ("List Bullet", "Ventures / LaunchPad data -- Team Quebec (LaunchPad)."),
        ("List Bullet", "Messaging and notifications -- Team Quebec (ConnectHub)."),
        ("Heading 3", "Strategy Pattern -- Ready for Swap"),
        ("Normal", "Delta uses the Strategy pattern for all external dependencies. When your real service is ready, swap the binding in AppServiceProvider -- no controller code changes needed:"),
        ("List Bullet", "CourseDataProviderInterface: MockCourseDataProvider (5 courses) -> real HTTP client to Papa (Experience Service)"),
        ("List Bullet", "CredentialDataProviderInterface: MockCredentialDataProvider (3 credentials) -> real HTTP client to Karl (Enrolment + Dashboard)"),
        ("List Bullet", "StudentProgressProviderInterface: MockStudentProgressProvider (random metrics) -> real aggregation from Papa + Karl (Dashboard)"),
    ]
    for style_name, text in reversed(sections):
        doc.paragraphs[integ_idx].insert_paragraph_before(text, style_name)

# ============================================================
# 10. Add Users Table Contract + Quebec questions
# ============================================================
error_fmt_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Error Format" and p.style.name == "Heading 2":
        error_fmt_idx = i
        break

if error_fmt_idx:
    additions = [
        ("Heading 4", "2. Users Table Contract"),
        ("Normal", "Delta seeds the users table identically across all three services. The minimum columns Delta requires:"),
        ("List Bullet", "id (bigint PK, not null) -- Auto-increment"),
        ("List Bullet", "name (varchar, not null) -- Display name"),
        ("List Bullet", "email (varchar, not null) -- Unique login email"),
        ("List Bullet", "password (varchar, not null) -- Hashed password (Delta never reads this, but column must exist)"),
        ("List Bullet", "role (varchar, not null) -- One of: school_admin, school_teacher, student, parent"),
        ("List Bullet", "school_id (bigint FK, not null) -- References schools.id"),
        ("List Bullet", "grade (smallint, nullable) -- Student grade level (8-12). Null for non-students."),
        ("List Bullet", "parent_of (bigint FK, nullable) -- References users.id. D1 supports one child per parent. Multi-child requires Quebec to own a parent_student pivot table."),
        ("Heading 4", "3. Open Questions for Quebec"),
        ("List Bullet", "Who owns the users table in production? Delta currently seeds it. In production, Quebec should be the source of truth. Delta needs read access only."),
        ("List Bullet", "How are new schools created? Delta has a schools table but no admin endpoint for creating schools. Is this Quebec's responsibility?"),
        ("List Bullet", "Role assignment flow -- Does Quebec assign roles at signup, or can admins change roles later? Delta's middleware checks role on every request."),
    ]
    for style_name, text in reversed(additions):
        doc.paragraphs[error_fmt_idx].insert_paragraph_before(text, style_name)

# ============================================================
# Save
# ============================================================
doc.save("C:/claude_workspaces/hatchloom/hatchloom-delta/API-CONTRACT.docx")
print("API-CONTRACT.docx updated successfully.")
