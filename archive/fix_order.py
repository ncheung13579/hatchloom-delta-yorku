"""Fix the reversed Data Scoping section in API-CONTRACT.docx."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from lxml import etree
from docx.oxml.ns import qn

doc = Document('API-CONTRACT.docx')

# Paragraphs 29-36 are in reverse order. We need to reverse them.
# 36 [Heading 2] Student and Parent Data Scoping  → should be first
# 35 [Normal] When a student or parent...         → should be second
# 34 [List Bullet] Students can only...           → third
# 33 [List Bullet] Parents can only...            → fourth
# 32 [List Bullet] Controller-level guards...     → fifth
# 31 [List Bullet] School-wide dashboard...       → sixth
# 30 [List Bullet] The student drill-down...      → seventh
# 29 [Normal] Three-tier enforcement...           → eighth (last)

# Get the XML elements for paragraphs 29-36
elements = [doc.paragraphs[i]._element for i in range(29, 37)]
parent = elements[0].getparent()

# Remove them all from the document
for elem in elements:
    parent.remove(elem)

# Find where "Services Overview" heading is now (it shifted)
services_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Services Overview':
        services_idx = i
        break

if services_idx is not None:
    services_elem = doc.paragraphs[services_idx]._element
    # Insert them back in reversed order (which is the correct order)
    for elem in reversed(elements):
        services_elem.addprevious(elem)
    print("Reordered data scoping section")
else:
    print("ERROR: Could not find Services Overview")

# Verify
for i in range(28, 46):
    p = doc.paragraphs[i]
    print(f'{i:3d} [{p.style.name}] {p.text[:120]}')

doc.save('API-CONTRACT.docx')
