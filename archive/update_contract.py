"""Update API-CONTRACT.docx with workpack-resolved decisions."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from copy import deepcopy

doc = Document('API-CONTRACT.docx')

changes = []

# === 1. Update Authentication section ===
# Find paragraph 28 (integration dependency note about auth)
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Integration dependency: This is mock auth only. Real JWT/token-based'):
        old = p.text
        p.text = (
            "Integration dependency: This is mock auth only. Per the project workpacks, "
            "the production auth model uses session tokens (not JWT). Team Quebec's Auth Service "
            "creates sessions in a database table (sessions: id, user_id, token, expires_at). "
            "The API Gateway (Role A) validates session tokens on incoming requests and propagates "
            "the authenticated user's identity to downstream services. Delta's services will not "
            "need to call Quebec per-request — the gateway handles validation centrally. "
            "See the \"Integration Requests\" section below for details."
        )
        for run in p.runs:
            run.font.size = Pt(10)
        changes.append(f"Updated auth integration note (para {i})")
        break

# === 2. Update version line ===
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Current version: D2'):
        p.text = (
            "Current version: D2.1 (March 2026) -- updated with workpack-confirmed decisions "
            "on auth architecture, progress ownership, parent model, and data scoping."
        )
        changes.append(f"Updated version line (para {i})")
        break

# === 3. Update Data Ownership — Data Delta does NOT provide ===
for i, p in enumerate(doc.paragraphs):
    # Fix the JWT auth line
    if 'User accounts, login, JWT auth' in p.text:
        p.text = (
            "User accounts, login, session-token auth -- Team Quebec (Auth). "
            "Workpacks confirm: session-token model, not JWT. Quebec stores sessions in DB. "
            "API Gateway validates tokens centrally and propagates identity downstream."
        )
        changes.append(f"Updated auth ownership line (para {i})")

    # Fix the parent-child line
    if 'Parent-child relationships (multi-child)' in p.text:
        p.text = (
            "Parent-child relationships -- Team Quebec. "
            "Workpacks confirm: canonical schema uses parent_student_links table (parent_id, student_id), "
            "not a single FK. Delta's D1 parent_of column is a temporary simplification. "
            "Quebec owns the link table; Delta will query it when real auth is integrated."
        )
        changes.append(f"Updated parent-child line (para {i})")

    # Fix the progress line
    if 'Per-student course progress -- Team Papa' in p.text:
        p.text = (
            "Per-student course progress -- Team Papa. "
            "Workpacks confirm: Papa's Course Service scope explicitly includes "
            "\"course progress tracking\" and \"course progress calculation: blocks completed, "
            "nodes completed within a block, overall percentage.\" "
            "Papa also tracks card completions and submission status."
        )
        changes.append(f"Updated progress ownership line (para {i})")

# === 4. Update Team Quebec Integration Requests ===
for i, p in enumerate(doc.paragraphs):
    # Update the heading for Quebec auth
    if p.text == '1. Real JWT/Token-Based Authentication':
        p.text = '1. Real Session-Token Authentication (Confirmed by Workpacks)'
        changes.append(f"Updated Quebec auth heading (para {i})")

    # Update the description under Quebec auth
    if p.text.startswith('We need to replace our mock authentication with real token verification.'):
        p.text = (
            "We need to replace our mock authentication with real session-token verification. "
            "Per the project workpacks (Team Quebec Sub-Pack Q1), the auth model is: "
            "Quebec creates a session record on login and returns a session token. "
            "The API Gateway (Role A) validates session tokens on every incoming request "
            "and propagates the authenticated user identity to downstream services. "
            "Delta does not need to validate tokens directly — the gateway does it."
        )
        changes.append(f"Updated Quebec auth description (para {i})")

    # Update the mock auth description
    if p.text.startswith('What we currently mock: MockAuthMiddleware in each service'):
        p.text = (
            "What we currently mock: MockAuthMiddleware in each service "
            "(e.g., dashboard-service/app/Http/Middleware/MockAuthMiddleware.php). "
            "This middleware maps hardcoded bearer tokens to user records via a TOKEN_MAP array. "
            "In production, this will be replaced by reading the user identity from headers "
            "injected by the API Gateway after session-token validation."
        )
        changes.append(f"Updated mock auth description (para {i})")

    # Update the users table contract - parent_of line
    if p.text.startswith('parent_of (bigint FK, nullable)'):
        p.text = (
            "parent_of (bigint FK, nullable) -- References users.id. D1 temporary simplification "
            "supporting one child per parent. Workpacks confirm: canonical schema uses a separate "
            "parent_student_links table (parent_id, student_id) owned by Quebec. "
            "Delta will migrate to querying this table when real auth is integrated."
        )
        changes.append(f"Updated parent_of field description (para {i})")

# === 5. Update Open Questions for Quebec ===
for i, p in enumerate(doc.paragraphs):
    if p.text == '3. Open Questions for Quebec':
        p.text = '3. Open Questions for Quebec (Partially Resolved by Workpacks)'
        changes.append(f"Updated open questions heading (para {i})")

    if p.text.startswith('Who owns the users table in production?'):
        p.text = (
            "Who owns the users table in production? RESOLVED: Quebec owns it "
            "(workpack Role B schema, table ownership: \"Users & Identity: users, sessions, "
            "parent_student_links → Quebec (Auth)\"). Delta needs read access only."
        )
        changes.append(f"Resolved users table question (para {i})")

    if p.text.startswith('How are new schools created?'):
        p.text = (
            "How are new schools created? PARTIALLY RESOLVED: Role B (Karl) owns the "
            "schema-per-school isolation strategy and school onboarding migration. "
            "Specific mechanism not yet defined — it is one of Karl's deliverables."
        )
        changes.append(f"Updated school creation question (para {i})")

# === 6. Add Data Scoping section after Authentication ===
# Find the right insertion point (after auth section, before Services Overview)
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Services Overview':
        insert_idx = i
        break

if insert_idx:
    # We need to insert paragraphs before "Services Overview"
    # In python-docx, we insert by manipulating the XML
    from docx.oxml.ns import qn
    services_elem = doc.paragraphs[insert_idx]._element

    # Build the paragraphs to insert (in reverse order since we insert before)
    texts_to_insert = [
        ("Heading 2", "Student and Parent Data Scoping"),
        ("Normal", (
            "When a student or parent authenticates, Delta automatically restricts "
            "what data they can see. This scoping is NOT documented per-endpoint above "
            "because it applies uniformly. Integrating teams should be aware of this behavior:"
        )),
        ("List Bullet", (
            "Students can only access their own data. The Enrolment Service auto-injects "
            "a student_id filter based on the authenticated user's role. Students calling "
            "list endpoints (e.g., GET /enrolments) will only see their own enrolments."
        )),
        ("List Bullet", (
            "Parents can only access their linked child's data. The same auto-filter applies, "
            "using the parent_of relationship (D1: single FK; production: parent_student_links table)."
        )),
        ("List Bullet", (
            "Controller-level guards enforce drill-down scoping. For endpoints like "
            "GET /dashboard/students/{studentId}, the controller checks that a student can only "
            "view their own record (403 if studentId != authenticated user's id) and a parent "
            "can only view their linked child's record."
        )),
        ("List Bullet", (
            "School-wide dashboard endpoints (GET /dashboard, GET /dashboard/reporting/*, "
            "GET /dashboard/widgets/*) are restricted to school_admin and school_teacher roles only. "
            "Students and parents cannot access aggregated school data."
        )),
        ("List Bullet", (
            "The student drill-down (GET /dashboard/students/{studentId}) is accessible by all roles, "
            "but students and parents are restricted to their own/linked data by controller guards."
        )),
        ("Normal", (
            "Three-tier enforcement: (1) Route middleware restricts which roles can reach an endpoint. "
            "(2) Controller guards verify the specific user can access the specific resource. "
            "(3) Service layer auto-injects filters based on role. This means a student using "
            "test-student-token will get different results than an admin using test-admin-token "
            "on the same endpoint."
        )),
    ]

    # Insert in reverse order so they end up in correct order
    for style_name, text in reversed(texts_to_insert):
        new_p = deepcopy(doc.paragraphs[0]._element)  # clone a paragraph element
        new_p.clear()  # clear its content
        # Create the paragraph properly
        from lxml import etree
        new_p = etree.SubElement(services_elem.getparent(), qn('w:p'))
        services_elem.getparent().remove(new_p)
        services_elem.addprevious(new_p)

        # Add run with text
        r = etree.SubElement(new_p, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')

        # Set paragraph style
        pPr = etree.SubElement(new_p, qn('w:pPr'))
        new_p.insert(0, pPr)
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pStyle.set(qn('w:val'), style_name.replace(' ', ''))

    changes.append("Added Student and Parent Data Scoping section")

# === 7. Add note about canonical data hierarchy to Papa's Course Catalogue section ===
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('What we currently mock: MockCourseDataProvider in experience-service'):
        # Find the next paragraph and add a note after it
        old = p.text
        p.text = (
            "What we currently mock: MockCourseDataProvider in "
            "experience-service/app/Services/MockCourseDataProvider.php, "
            "which returns 5 hardcoded courses with 2 blocks each. "
            "Note: Delta's D1 mock uses a flat Course → Block structure with block fields "
            "(id, name, type). The canonical data hierarchy from Papa's workpack is deeper: "
            "Course → Block → Node → Activity Card (10 types). Block types in the mock "
            "(\"lesson\", \"challenge\") are D1 simplifications. When Papa's real API is available, "
            "the response shape for blocks will likely include nested nodes and activity cards."
        )
        changes.append(f"Updated course mock description with hierarchy note (para {i})")
        break

# === 8. Update Papa progress section with workpack confirmation ===
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('We need per-student progress information: courses completed'):
        p.text = (
            "We need per-student progress information: courses completed, courses in progress, "
            "completion percentage, and blocks completed. "
            "Workpacks confirm: Papa's Course Service scope explicitly includes "
            "\"course progress tracking\" and \"course progress calculation.\" "
            "Papa Sub-Pack P1 owns enrollment and progress; Sub-Pack P3 owns card completion tracking."
        )
        changes.append(f"Updated Papa progress description with workpack confirmation (para {i})")
        break

# === 9. Update dashboard endpoint auth descriptions ===
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Returns the full aggregated dashboard for the authenticated school admin.':
        p.text = (
            "Returns the full aggregated dashboard for the authenticated school admin. "
            "Restricted to school_admin and school_teacher roles only — students and parents "
            "cannot access this endpoint (returns 403)."
        )
        changes.append(f"Updated dashboard overview access note (para {i})")
        break

doc.save('API-CONTRACT.docx')

print(f"Applied {len(changes)} changes:")
for c in changes:
    print(f"  - {c}")
